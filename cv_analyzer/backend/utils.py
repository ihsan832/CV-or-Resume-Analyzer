"""
Text extraction utilities for PDF and DOCX files.
"""

import io


def extract_text_from_file(file_bytes: bytes, extension: str) -> str:
    """
    Extract plain text from a PDF or DOCX file.

    Args:
        file_bytes: Raw file bytes.
        extension:  File extension, e.g. '.pdf' or '.docx'.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the extension is unsupported.
        RuntimeError: If extraction fails.
    """
    ext = extension.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise RuntimeError(f"PDF extraction error: {e}")

    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    text_parts = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction error: {e}")

    return "\n".join(text_parts)
